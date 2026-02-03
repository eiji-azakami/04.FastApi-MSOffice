"""
Word出力API

このモジュールは、Wordファイル出力のサンプルプログラムです。
"""

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from urllib.parse import quote

router = APIRouter(tags=["word"])


@router.get("/word/export")
def export_word():
    """
    Wordファイル出力
    """
    # doc生成
    document = Document()

    # タイトル部分
    title = document.add_heading("委 任 状", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.size = Pt(24)

    # 受任者情報
    paragraph = document.add_paragraph("受任者")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph("　住所 （受任者の住所）").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph("　氏名 （受任者の氏名）").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 本文
    paragraph = document.add_paragraph("私は、上記のものを代理人と定め、下記の権限を委任します。")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 記
    document.add_paragraph("")
    document.add_paragraph("記").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph("")

    # 権限内容
    document.add_paragraph(
        "　次にあげる納税証明書の申請及び受領に関する権限"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph("　　1. ○○○○").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph("　　2. ○○○○").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph("")
    document.add_paragraph("")

    # 日付と署名欄
    document.add_paragraph("○○○○年 ○○月 ○○日").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph(
        "　　　　　　　　　　　　　　　　　　　　　受任者"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph(
        "　　　　　　　　　　　　　　　　　　　　　　住所 （受任者の住所）"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph(
        "　　　　　　　　　　　　　　　　　　　　　　氏名 （受任者の氏名）"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph(
        "　　　　　　　　　　　　　　　　　　　　電話番号 999-999-9999"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph(
        "　　　　　　　　　　　　　　　　　　　　生年月日 ○○○○年 ○○月 ○○日"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # メモリ内でファイルを保存
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # ファイル名をURLエンコード
    encoded_filename = quote("委任状.docx")

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )
